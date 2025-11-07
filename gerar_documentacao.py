from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

print("Gerando documentação...")

# Criar documento
doc = Document()

# ============================================================================
# CAPA
# ============================================================================
# Título principal
titulo = doc.add_heading('Sistema de Gerenciamento', 0)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
titulo_run = titulo.runs[0]
titulo_run.font.size = Pt(28)
titulo_run.font.color.rgb = RGBColor(0, 51, 102)

# Subtítulo
subtitulo = doc.add_heading('Clínica Veterinária', 1)
subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitulo_run = subtitulo.runs[0]
subtitulo_run.font.size = Pt(22)
subtitulo_run.font.color.rgb = RGBColor(102, 51, 153)

doc.add_paragraph('\n' * 3)

# Informações do projeto
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = info.add_run('Projeto Integrador Extensionista 3\n')
info_run.font.size = Pt(14)
info_run = info.add_run('Universidade de Marília - UNIMAR\n\n')
info_run.font.size = Pt(14)
info_run = info.add_run('Desenvolvido por: Marcio Santos\n')
info_run.font.size = Pt(14)
info_run.bold = True
info_run = info.add_run(f'RA: 13119972\n\n')
info_run.font.size = Pt(14)
info_run.bold = True
info_run = info.add_run(f'Novembro de 2025\n')
info_run.font.size = Pt(12)

doc.add_page_break()

# ============================================================================
# SUMÁRIO (Manual)
# ============================================================================
doc.add_heading('Sumário', 1)
sumario_items = [
    '1. Descrição do Sistema',
    '2. Tecnologias Utilizadas',
    '3. Funcionalidades Implementadas',
    '4. Instalação e Execução',
    '5. Estrutura do Projeto',
    '6. Prints das Telas',
    '7. Conclusão'
]
for item in sumario_items:
    p = doc.add_paragraph(item)
    p.style = 'List Number'

doc.add_page_break()

# ============================================================================
# 1. DESCRIÇÃO DO SISTEMA
# ============================================================================
doc.add_heading('1. Descrição do Sistema', 1)
doc.add_paragraph(
    'O Sistema de Gerenciamento da Clínica Veterinária Unimar é uma aplicação web completa '
    'desenvolvida para facilitar a administração de clínicas veterinárias. O sistema permite '
    'o gerenciamento integrado de donos de pets, animais de estimação, veterinários e '
    'agendamento de consultas.'
)
doc.add_paragraph(
    'A solução foi desenvolvida com foco em usabilidade, permitindo que qualquer usuário, '
    'mesmo sem conhecimentos técnicos avançados, possa realizar cadastros, consultas, '
    'atualizações e exclusões de registros de forma intuitiva.'
)

# Objetivos
doc.add_heading('Objetivos do Sistema:', 2)
objetivos = [
    'Centralizar informações de donos, pets e veterinários',
    'Facilitar o agendamento e controle de consultas',
    'Proporcionar interface amigável e responsiva',
    'Garantir integridade e persistência dos dados',
    'Otimizar o fluxo de trabalho da clínica veterinária'
]
for obj in objetivos:
    doc.add_paragraph(obj, style='List Bullet')

doc.add_page_break()

# ============================================================================
# 2. TECNOLOGIAS UTILIZADAS
# ============================================================================
doc.add_heading('2. Tecnologias Utilizadas', 1)

doc.add_heading('Flask (Back-end)', 2)
doc.add_paragraph(
    'Framework web em Python para criar aplicações web e APIs. É leve, simples e flexível. '
    'No projeto, o Flask é responsável por todo o back-end, gerenciando o banco de dados '
    'e servindo as páginas HTML através de rotas definidas no arquivo backend/app.py.'
)

doc.add_heading('Bootstrap (Front-end)', 2)
doc.add_paragraph(
    'Framework CSS front-end para criar interfaces responsivas e bonitas rapidamente. '
    'Utilizado em todas as páginas HTML para estilização, incluindo botões, cards, '
    'barra de navegação e formulários. Permite que o sistema seja acessado de qualquer '
    'dispositivo (desktop, tablet ou smartphone) com layout adaptado.'
)

doc.add_heading('SQLite (Banco de Dados)', 2)
doc.add_paragraph(
    'Sistema de gerenciamento de banco de dados relacional leve e embutido. '
    'Perfeito para aplicações de pequeno a médio porte, não requer instalação '
    'de servidor separado. Armazena todos os dados em um único arquivo (clinica.db).'
)

doc.add_heading('Python', 2)
doc.add_paragraph(
    'Linguagem de programação de alto nível, utilizada tanto no back-end (Flask) '
    'quanto nos scripts auxiliares do projeto.'
)

doc.add_heading('HTML5 e JavaScript', 2)
doc.add_paragraph(
    'HTML5 para estruturação das páginas web e JavaScript para interatividade '
    'e comunicação assíncrona com a API REST do back-end.'
)

doc.add_page_break()

# ============================================================================
# 3. FUNCIONALIDADES IMPLEMENTADAS
# ============================================================================
doc.add_heading('3. Funcionalidades Implementadas', 1)

doc.add_heading('3.1 Módulo de Donos de Pets', 2)
funcionalidades_donos = [
    'Cadastro completo com nome, telefone, email, endereço e CEP',
    'Listagem de todos os donos cadastrados',
    'Busca e filtros na tabela',
    'Edição de informações existentes',
    'Exclusão de registros',
    'Máscaras de entrada para telefone (formato brasileiro) e CEP',
    'Validação de dados no formulário'
]
for func in funcionalidades_donos:
    doc.add_paragraph(func, style='List Bullet')

doc.add_heading('3.2 Módulo de Pets (Animais de Estimação)', 2)
funcionalidades_pets = [
    'Cadastro com nome, espécie, raça, idade e peso',
    'Vinculação automática com o dono do pet',
    'Listagem completa com informações do dono',
    'Visualização detalhada de cada animal',
    'Edição de dados cadastrais',
    'Exclusão de registros',
    'Interface intuitiva com seleção do dono via dropdown'
]
for func in funcionalidades_pets:
    doc.add_paragraph(func, style='List Bullet')

doc.add_heading('3.3 Módulo de Veterinários', 2)
funcionalidades_vets = [
    'Cadastro com nome, CRMV, especialidade, email e telefone',
    'Validação de CRMV único (não permite duplicatas)',
    'Listagem de todos os profissionais',
    'Edição de informações profissionais',
    'Exclusão de registros',
    'Máscara de entrada para CRMV (formato: 00000-UF)',
    'Gestão de especialidades médicas veterinárias'
]
for func in funcionalidades_vets:
    doc.add_paragraph(func, style='List Bullet')

doc.add_heading('3.4 Módulo de Agendamento de Consultas', 2)
funcionalidades_consultas = [
    'Agendamento de consultas com data e hora',
    'Seleção do pet e veterinário',
    'Tipos de atendimento: consulta, vacina, cirurgia, exame',
    'Campo para motivo/observações',
    'Controle de status (agendada, realizada, cancelada)',
    'Listagem completa com informações integradas',
    'Visualização de histórico de consultas'
]
for func in funcionalidades_consultas:
    doc.add_paragraph(func, style='List Bullet')

doc.add_page_break()

# ============================================================================
# 4. INSTALAÇÃO E EXECUÇÃO
# ============================================================================
doc.add_heading('4. Instalação e Execução', 1)

doc.add_heading('4.1 Requisitos', 2)
requisitos = [
    'Python 3.8 ou superior',
    'Navegador web moderno (Chrome, Firefox, Edge)'
]
for req in requisitos:
    doc.add_paragraph(req, style='List Bullet')

doc.add_heading('4.2 Como Executar', 2)
doc.add_paragraph('1. Descompacte o arquivo em uma pasta')
doc.add_paragraph('2. Duplo clique em: INICIAR_SERVIDOR.bat')
doc.add_paragraph('3. O navegador abrirá automaticamente')
doc.add_paragraph('4. Sistema disponível em: http://127.0.0.1:5000')

doc.add_paragraph('')
doc.add_paragraph('Problemas? Consulte o arquivo: docs/SOLUCAO_PROBLEMAS.md')

doc.add_page_break()

# ============================================================================
# 5. ESTRUTURA DO PROJETO
# ============================================================================
doc.add_heading('5. Estrutura do Projeto', 1)

doc.add_paragraph('O projeto está organizado da seguinte forma:')
doc.add_paragraph('')

estrutura = """RA13119972-23/
├── backend/
│   └── app.py              # API Flask
├── database/
│   └── schema.sql          # Estrutura do banco
├── frontend/
│   ├── index.html          # Página inicial
│   ├── donos.html          # Donos
│   ├── pets.html           # Pets
│   ├── veterinarios.html   # Veterinários
│   └── consultas.html      # Consultas
├── INICIAR_SERVIDOR.bat    # Inicialização (recomendado)
├── INICIAR.py              # Script Python
└── requirements.txt        # Dependências"""

p = doc.add_paragraph(estrutura)
p.style = 'Intense Quote'

doc.add_page_break()

# ============================================================================
# 6. PRINTS DAS TELAS
# ============================================================================
doc.add_heading('6. Prints das Telas do Sistema', 1)

doc.add_paragraph(
    'IMPORTANTE: Tire prints do sistema em funcionamento e insira nas páginas seguintes. '
    'Os prints devem mostrar as telas com dados preenchidos.'
)
doc.add_paragraph('')

# Seções para prints
prints_secoes = [
    ('6.1 Página Inicial', 
     'Tela inicial mostrando os quatro cards principais.'),
    
    ('6.2 Cadastro de Donos', 
     'Tela de cadastro e listagem de donos.'),
    
    ('6.3 Cadastro de Pets', 
     'Tela de cadastro e listagem de pets.'),
    
    ('6.4 Cadastro de Veterinários', 
     'Tela de cadastro e listagem de veterinários.'),
    
    ('6.5 Agendamento de Consultas', 
     'Tela de agendamento e listagem de consultas.')
]

for titulo, descricao in prints_secoes:
    doc.add_heading(titulo, 2)
    doc.add_paragraph(descricao)
    doc.add_paragraph('')
    doc.add_paragraph('[INSERIR PRINT AQUI]')
    doc.add_paragraph('')
    doc.add_paragraph('_' * 80)
    doc.add_paragraph('')

doc.add_page_break()

# ============================================================================
# 7. CONCLUSÃO
# ============================================================================
doc.add_heading('7. Conclusão', 1)
doc.add_paragraph(
    'O Sistema de Gerenciamento da Clínica Veterinária foi desenvolvido com sucesso, '
    'atendendo aos requisitos propostos. O sistema apresenta interface intuitiva, '
    'funcionalidades completas de CRUD para todos os módulos e integração eficiente '
    'entre front-end, back-end e banco de dados.'
)
doc.add_paragraph(
    'A aplicação está pronta para uso em clínicas veterinárias, proporcionando '
    'uma solução completa para gerenciamento de informações.'
)

doc.add_paragraph('')
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('___________________________________\n')
run = p.add_run('Marcio Santos\n')
run.bold = True
run = p.add_run('RA: 13119972\n')
run = p.add_run('UNIMAR - 2025')

# Salvar documento
output_file = 'DOCUMENTACAO_RA13119972.docx'
doc.save(output_file)
print("✅ Documentação criada com sucesso!")
print(f"📄 Arquivo: {output_file}")
print("\n⚠️ IMPORTANTE:")
print("1. Tire prints das telas do sistema funcionando")
print("2. Abra o arquivo DOCX e insira os prints na seção 6")
print("3. Salve e converta para PDF se necessário")
